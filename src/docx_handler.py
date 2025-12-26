import os
from docx import Document

def extract_text_from_docx(file_path):
    """
    Extracts all text from a .docx file (paragraphs and tables) 
    to send to Gemini.
    Returns: A list of non-empty strings.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)
    full_text = []

    # 1. Extract from Standard Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # 2. Extract from Tables (Crucial for Resumes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)

    return full_text

def replace_text_preserving_runs(paragraph, search_text, new_text):
    """
    Your custom logic: Replaces text inside a paragraph while keeping formatting.
    Returns True if a replacement was made.
    """
    if search_text not in paragraph.text:
        return False

    # Loop to handle multiple occurrences in the same paragraph
    while search_text in paragraph.text:
        text = paragraph.text
        match_start = text.find(search_text)
        match_end = match_start + len(search_text)
        
        current_pos = 0
        replacement_inserted = False

        for run in paragraph.runs:
            run_len = len(run.text)
            run_start = current_pos
            run_end = current_pos + run_len
            
            # Logic: overlap exists if (StartA <= EndB) and (EndA >= StartB)
            if run_start < match_end and run_end > match_start:
                
                # Calculate which part of this run to keep
                keep_before = ""
                if match_start > run_start:
                    keep_before = run.text[:match_start - run_start]
                
                keep_after = ""
                if match_end < run_end:
                    keep_after = run.text[match_end - run_start:]

                # Perform the Swap
                if not replacement_inserted:
                    run.text = keep_before + new_text + keep_after
                    replacement_inserted = True
                else:
                    # If new text was already inserted, this run is part of the old text to be removed
                    run.text = keep_after

            current_pos += run_len
        
        # Break to avoid infinite loops if replacement contains search text
        break
        
    return True

def process_resume_updates(input_path, output_path, replacements_list):
    """
    The Main Worker Function.
    Args:
        input_path (str): Path to original resume.
        output_path (str): Path to save updated resume.
        replacements_list (list): List of dicts [{'original': '...', 'new': '...'}, ...]
    """
    try:
        doc = Document(input_path)
        replacements_count = 0

        # Iterate through every replacement request from Gemini
        for item in replacements_list:
            # --- CHANGE THESE TWO LINES ---
            search_text = item['original_exact_text']
            new_text = item['replacement_text']
            # -----------------------------

            # A. Scan Paragraphs
            for para in doc.paragraphs:
                # We need to handle DELETE operations where the text is empty
                if search_text and replace_text_preserving_runs(para, search_text, new_text):
                    replacements_count += 1

            # B. Scan Tables (Resumes often have layout tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if replace_text_preserving_runs(para, search_text, new_text):
                                replacements_count += 1

        doc.save(output_path)
        print(f"✓ Saved updated resume to: {output_path}")
        print(f"✓ Total text replacements made: {replacements_count}")
        return True

    except Exception as e:
        print(f"❌ Error processing resume: {e}")
        return False